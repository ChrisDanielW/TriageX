# generate_testing_tables_docx.py
# Requires: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)

def add_table(doc, headers, rows, style="Light List Accent 1"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
    # Data rows
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")  # spacing
    return table

def main(output_path="TriageX_Test_Tables.docx"):
    doc = Document()
    add_heading(doc, "Testing Tables", level=1)
    add_paragraph(doc, "Prepared templates for Unit, Integration, End‑to‑End, Performance, Security, Accessibility, and Resilience testing. Fill Actual Output and Status after execution.")

    # Common headers
    headers = ["Test Case ID", "Test Case Description", "Input", "Expected Output", "Actual Output", "Status"]

    # Tables data
    tables = [
        {
            "label": "Table UT-1: Unit Tests — Front-End Components",
            "desc": "Verifies component rendering, validation, and UI logic in React/TypeScript.",
            "rows": [
                ["UT001", "[UI] LoginForm email validation", "Enter invalid email, blur field", "Inline error shown; submit disabled", "—", "TBD"],
                ["UT002", "[UI] HospitalCombobox filtering", "Type “Baptist” in search", "List filtered to hospitals matching “Baptist”", "—", "TBD"],
                ["UT003", "[UI] PatientCard severity badge", "triage_level=CRITICAL", "Red CRITICAL chip rendered", "—", "TBD"],
                ["UT004", "[UI] CriticalityProgressBar bounds", "score=11", "Value clamped to 10 with full bar", "—", "TBD"],
                ["UT005", "[Logic] play/stop emergency sound", "Trigger play, then stop", "Audio loops until stop; stop silences audio", "—", "TBD"],
                ["UT006", "[Auth] RequireRole gate", "Render child with missing role", "Child not rendered; redirect or notice", "—", "TBD"],
                ["UT007", "[Util] ETA formatter", "420 seconds", "“7 min” string returned", "—", "TBD"],
                ["UT008", "[Hook] use-toast stacking", "Fire 3 toasts quickly", "All toasts visible; auto-dismiss on timers", "—", "TBD"],
            ],
        },
        {
            "label": "Table UT-2: Unit Tests — Back-End/Service Layer",
            "desc": "Validates server utilities/controllers and schemas used by REST/Socket services.",
            "rows": [
                ["UT101", "[Auth] Login rejects bad password", "Valid email, wrong password", "401 Unauthorized", "—", "TBD"],
                ["UT102", "[Schema] enroutealert validation", "Missing priority", "Validation error “priority required”", "—", "TBD"],
                ["UT103", "[State] status transition rules", "accepted → accepted", "Idempotent; single state, no duplicate", "—", "TBD"],
                ["UT104", "[Socket] room naming", "hospital_id=abc joins", "Client joins room “hospital:abc”", "—", "TBD"],
                ["UT105", "[Util] severity mapping", "Vitals imply critical", "Returns “CRITICAL” label", "—", "TBD"],
            ],
        },
        {
            "label": "Table IT-1: Integration Tests — Auth + Session",
            "desc": "Ensures client and API work together for login, session persistence, and guards.",
            "rows": [
                ["IT001", "[Auth] Email/password login flow", "Submit valid creds", "Session cookie set; redirect to dashboard", "—", "TBD"],
                ["IT002", "[Auth] /auth/me after refresh", "Reload dashboard", "User + hospital context restored", "—", "TBD"],
                ["IT003", "[Guard] Protected route access", "Visit /hospital without session", "Redirect to login", "—", "TBD"],
                ["IT004", "[Logout] Session clear", "Click logout", "Cookie cleared; back to login", "—", "TBD"],
            ],
        },
        {
            "label": "Table IT-2: Integration Tests — Real-Time Alerts (Socket + UI)",
            "desc": "Validates socket-driven updates appear correctly in the dashboard and modals.",
            "rows": [
                ["IT101", "[Alerts] New alert renders card", "Emit alert:new", "New PatientCard appears with details", "—", "TBD"],
                ["IT102", "[Alerts] ETA updates in place", "Emit alert:update (eta=7m)", "Card ETA updates to “7 min”", "—", "TBD"],
                ["IT103", "[Audio] Critical tone plays", "Emit critical alert", "Looping alert sound starts", "—", "TBD"],
                ["IT104", "[Action] Accept stops audio", "Click Accept in modal", "Status→accepted; audio stops", "—", "TBD"],
                ["IT105", "[Action] Reject removes card", "Click Reject", "Card removed/hidden; status persisted", "—", "TBD"],
            ],
        },
        {
            "label": "Table IT-3: Integration Tests — Data Linking",
            "desc": "Confirms en-route snapshots connect to patient records and related collections.",
            "rows": [
                ["IT201", "[Link] Alert→Patient record", "Mark Arrived", "Patient created/linked to alert patient_id", "—", "TBD"],
                ["IT202", "[Vitals] Time-series append", "Send new vitals sample", "Latest vitals visible in details timeline", "—", "TBD"],
                ["IT203", "[Triage] Assessment attach", "Submit triage form", "triageassessment saved with user_id", "—", "TBD"],
            ],
        },
        {
            "label": "Table E2E-1: End-to-End Flows",
            "desc": "Validates key user journeys across the stack from login to decisions.",
            "rows": [
                ["E2E001", "Login → Accept critical alert", "Valid creds; open first critical", "Case accepted in ≤2 clicks; audio silenced", "—", "TBD"],
                ["E2E002", "Critical→Arrived lifecycle", "Accept then mark arrived", "Dashboard reflects arrived; history updated", "—", "TBD"],
                ["E2E003", "Reject with reason (if enabled)", "Reject case", "Case hidden; reason stored/shown", "—", "TBD"],
            ],
        },
        {
            "label": "Table PT-1: Performance/Load Tests",
            "desc": "Measures responsiveness under realistic and peak conditions.",
            "rows": [
                ["PT001", "Dashboard render with 100 cards", "Seed 100 alerts", "First paint < 2s; main thread idle > 60%", "—", "TBD"],
                ["PT002", "Socket burst handling", "50 updates in 5s", "No dropped updates; UI stutter < 200ms", "—", "TBD"],
                ["PT003", "Modal open latency", "Click any card", "Modal content interactive < 150ms", "—", "TBD"],
                ["PT004", "Audio start latency", "Receive critical alert", "Sound begins < 300ms from event", "—", "TBD"],
            ],
        },
        {
            "label": "Table ST-1: Security Tests",
            "desc": "Confirms access control, cookie security, and input sanitization.",
            "rows": [
                ["ST001", "Block unauthenticated access", "Direct /hospital URL", "302 to /login", "—", "TBD"],
                ["ST002", "RBAC on sensitive actions", "User without role attempts Accept", "403/blocked UI; no state change", "—", "TBD"],
                ["ST003", "Cookie flags", "Inspect session cookie", "HttpOnly, Secure, SameSite set", "—", "TBD"],
                ["ST004", "CORS policy", "Call API from untrusted origin", "Request blocked", "—", "TBD"],
                ["ST005", "XSS in fields", "Inject <script> in notes", "Render escaped; no script executes", "—", "TBD"],
            ],
        },
        {
            "label": "Table AT-1: Accessibility/UX Tests",
            "desc": "Ensures the UI is operable and perceivable for all users.",
            "rows": [
                ["AX001", "Keyboard-only navigation", "Tab through Login and modal", "Reachable controls; visible focus", "—", "TBD"],
                ["AX002", "ARIA and labels", "Inspect Patient Details", "Landmarks/labels present and meaningful", "—", "TBD"],
                ["AX003", "Contrast of severity chips", "Check CRITICAL badge", "Contrast ≥ 4.5:1", "—", "TBD"],
                ["AX004", "Live region for alerts", "Receive alert:new", "Screen reader announces new alert", "—", "TBD"],
            ],
        },
        {
            "label": "Table RT-1: Resilience/Recovery Tests",
            "desc": "Validates stability during faults and network changes.",
            "rows": [
                ["RS001", "Socket reconnect", "Drop network 10s, restore", "Auto-reconnect; no duplicate cards", "—", "TBD"],
                ["RS002", "Idempotent Accept", "Double-click Accept", "Single server transition; UI consistent", "—", "TBD"],
                ["RS003", "Offline fallback", "Go offline on dashboard", "Graceful message; retries/poll on return", "—", "TBD"],
                ["RS004", "Error toast surface", "Force 500 on accept", "Non-blocking error toast; no state drift", "—", "TBD"],
            ],
        },
    ]

    # Build document
    for t in tables:
        add_heading(doc, t["label"], level=2)
        add_paragraph(doc, f"Description: {t['desc']}")
        add_table(doc, headers, t["rows"])
        doc.add_paragraph("")  # spacing

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    main()